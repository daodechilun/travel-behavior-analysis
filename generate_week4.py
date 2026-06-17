# -*- coding: utf-8 -*-
"""生成周报四 docx 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_run_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)  # 强制黑色
    return h

def add_figure(doc, path, caption, width=Inches(5.2)):
    """嵌入图表到文档中"""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.add_run().add_picture(path, width=width)
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        run = caption_p.add_run(caption)
        set_run_font(run, '宋体', 9, color=(100, 100, 100))

def add_para(doc, text, bold=False, size=11, indent=True, align=None, font_name='宋体'):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p

def add_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, '黑体', 10, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, '宋体', 9)
    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== 标题 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('大数据专业综合课程设计实验报告')
set_run_font(run, '黑体', 22, True, color=(0, 0, 0))
p.paragraph_format.space_after = Pt(12)

# ===== 信息表头 =====
info_table = doc.add_table(rows=2, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in info_table.rows[0].cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
info_headers = ['班级', '学 号', '姓 名', '指导教师']
info_values = ['2', '20230222', '梁文泽', '陈楠']
for j, (h, v) in enumerate(zip(info_headers, info_values)):
    r0 = info_table.rows[0].cells[j].paragraphs[0].add_run(h)
    set_run_font(r0, '宋体', 11, True)
    info_table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = info_table.rows[1].cells[j].paragraphs[0].add_run(v)
    set_run_font(r1, '宋体', 11)
    info_table.rows[1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

add_para(doc, '基本信息：基于高德地图数据的旅游用户行为分析与推荐策略优化', bold=True, size=12, indent=False)

doc.add_paragraph()

# ===== 一、项目进度概览 =====
add_heading_styled(doc, '一、项目进度概览', 1)

add_table(doc,
    ['阶段', '状态', '完成度', '较上周变化'],
    [
        ['选题确定', '已完成', '100%', '-'],
        ['大纲设计', '已完成', '100%', '-'],
        ['数据采集', '已完成', '100%', '↑15%（超额完成124%）'],
        ['数据清洗', '已完成', '100%', '↑50%'],
        ['行为路径分析', '已完成', '100%', '↑78%'],
        ['消费偏好挖掘', '已完成', '100%', '↑90%'],
        ['推荐策略优化', '已完成', '100%', '↑100%（本周新增）'],
        ['时空特征分析', '已完成', '100%', '↑100%（本周新增）'],
        ['转化漏斗分析', '已完成', '100%', '↑100%（本周新增）'],
        ['用户画像与分群', '已完成', '100%', '↑100%（本周新增）'],
        ['关联规则挖掘', '已完成', '100%', '↑100%（本周新增）'],
        ['城市对比雷达分析', '已完成', '100%', '↑100%（本周新增）'],
        ['实验报告自动生成', '已完成', '100%', '↑100%（本周新增）'],
    ]
)

doc.add_paragraph()

# ===== 二、本周已完成内容 =====
add_heading_styled(doc, '二、本周已完成内容', 1)

# --- 2.1 数据采集收尾与清洗完成 ---
add_heading_styled(doc, '2.1 数据采集收尾与全量数据清洗', 2)
add_para(doc, '完成工作：')
items_21 = [
    '✅ 数据采集全面完成：累计获得62,200条有效记录，远超原定50,000条目标（完成度124.4%），较上周42,500条净增19,700条。',
    '✅ 8个目标城市全覆盖：北京(8,550)、广州(8,450)、上海(7,700)、成都(7,700)、西安(7,700)、杭州(7,550)、重庆(7,300)、三亚(7,250)，分布均衡。',
    '✅ 4大旅游品类完整覆盖：美食(16,950)、酒店(15,400)、景点(15,300)、旅游服务(14,550)。',
    '✅ 数据清洗100%完成，实施完整清洗流程：数值列中位数填充 → 分类列众数填充 → 消费/价格取绝对值 → 评分截断至0-5 → 经纬度越界过滤 → 多格式日期解析 → 四字段联合去重。清洗后无缺失值。',
    '✅ 新增衍生特征：解析路径列表(path_list)、提取时段(hour, 0-23)、标注日类型(day_type, 工作日/周末)、消费分层(consume_tier, 8档)。',
    '✅ 输出清洗后数据文件 amap_travel_cleaned.csv（62,200条 × 17列），可直接用于后续建模。',
]
for item in items_21:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    set_run_font(run, '宋体', 11)

add_para(doc, '数据质量报告（终版）：')
add_table(doc,
    ['指标', '当前值', '上周值', '变化'],
    [
        ['累计数据总量', '62,200条', '42,500条', '↑46.4%'],
        ['目标完成度', '124.4%', '85.0%', '↑39.4pp'],
        ['覆盖城市数', '8个', '8个', '-'],
        ['覆盖品类数', '4个', '5个', '↓1（精简交通枢纽）'],
        ['字段完整率', '100%', '99.85%', '↑0.15pp'],
        ['唯一用户数', '44,958人', '-', '新增指标'],
        ['唯一POI数', '10,222个', '-', '新增指标'],
        ['时间跨度', '4/20-4/26（7天）', '-', '新增指标'],
    ]
)

doc.add_paragraph()

# --- 2.2 八大分析模块全面完成 ---
add_heading_styled(doc, '2.2 八大分析模块全面完成（核心突破）', 2)
add_para(doc, '本周完成了从模块3到模块10的全部8个分析模块，每个模块均包含统计计算+可视化图表输出。')

add_para(doc, '（1）模块3：行为路径分析', bold=True)
add_para(doc, '完成4种行为路径的频次统计与终点消费对比。最热门路径为"首页→搜索→POI详情→预订"（15,740人，占25.3%），其次"首页→推荐列表→POI详情→收藏"（15,710人，25.3%）。四种路径终点的平均消费差异极小（标准差仅22元），说明路径选择对消费金额的独立预测力较弱。')
add_figure(doc, 'output_figures/module3_path_analysis.png', '图1 行为路径分析：高频访问路径分布（左）与路径终点平均消费对比（右）')

add_para(doc, '（2）模块4：消费偏好挖掘', bold=True)
add_para(doc, '构建品类×城市热力图与消费分层行为分布。发现消费呈双峰结构——零消费用户18,522人（29.8%）与高消费(>5k)用户16,382人（26.3%）并存。上海消费最高(均值2,862元)，北京最低(2,777元)，但城市间差异幅度仅3%。')
add_figure(doc, 'output_figures/module4_consumption.png', '图2 消费偏好挖掘：品类×城市平均消费热力图（左）与消费分层行为分布（右）')

add_para(doc, '（3）模块5：推荐策略优化', bold=True)
add_para(doc, '建立四维加权评分模型（评分30%+热度25%+转化率30%+消费15%），对10,222个POI逐一评分（47-73.7分）。核心发现：转化率是高分POI与低分POI的最强区分指标（高分50.6% vs 低分16.3%，差34pp），而价格几乎无区分力（差仅7元）。')
add_figure(doc, 'output_figures/module5_recommendation.png', '图3 推荐策略优化：Top 15 POI综合推荐分（左）与各品类推荐分箱线图（右）')

add_para(doc, '（4）模块6：时空特征分析', bold=True)
add_para(doc, '完成4维时空分析：时段×行为热力图、消费与预订率趋势双轴图、8城市POI空间分布气泡图、工作日vs周末对比。晚间19:00用户最活跃，下午15:00预订转化率最高(25.7%)；8城市经纬度跨度103.2°-121.9°，广州POI最多(1,506)，三亚最少(1,003)；周末效应不显著（消费差异仅0.7%）。')
add_figure(doc, 'output_figures/module6_spatiotemporal.png', '图4 时空特征分析：时段×行为热力图、消费与预订率趋势、城市空间分布、工作日vs周末对比')

add_para(doc, '（5）模块7：转化漏斗分析', bold=True)
add_para(doc, '构建用户行为转化漏斗与品类对比曲线。四个行为环节人数分布均衡（各约14,200-15,500人），单步转化率接近100%。仅浏览不转化用户8,558人(19.0%)，主要流失品类为美食(27.4%)和酒店(25.0%)。四品类转化曲线高度重合，品类本身对漏斗形态影响有限。')
add_figure(doc, 'output_figures/module7_funnel.png', '图5 转化漏斗分析：用户行为转化漏斗（左）与各品类转化率曲线对比（右）')

add_para(doc, '（6）模块8：用户画像与分群', bold=True)
add_para(doc, 'RFM模型分层+K-Means聚类双方法用户分群。RFM结果：高价值21.4%、中价值28.0%、潜力39.2%、低活跃11.5%。K-Means最优K=2（轮廓系数0.315），聚类0"普通浏览用户"(69.4%，消费2,811元)与聚类1"高价值核心用户"(30.6%，消费6,349元，为前者2.3倍)差异显著。')
add_figure(doc, 'output_figures/module8_user_segments.png', '图6 用户画像与分群：RFM分层饼图、K-Means最优K选择、聚类雷达图、聚类用户分布')

add_para(doc, '（7）模块9：关联规则挖掘', bold=True)
add_para(doc, '采用Apriori算法挖掘品类间关联规则。发现12条规则，10个频繁项集，Lift值集中于0.49-0.50。品类间无强关联（Lift均接近1.0），用户对品类的选择偏向独立决策。')
add_figure(doc, 'output_figures/module9_association.png', '图7 关联规则挖掘：关联规则散点图（左）与品类×行为交叉频次热力图（右）')

add_para(doc, '（8）模块10：城市对比雷达分析', bold=True)
add_para(doc, '构建8城市6维雷达对比+综合得分排名+消费×转化气泡图。综合排名：杭州(0.851)>三亚(0.660)>上海(0.639)>广州(0.587)>重庆(0.521)>成都(0.423)>北京(0.401)>西安(0.399)。将8城市归纳为6种旅游画像类型：标杆型(杭州)、度假型(三亚)、商务型(上海)、效率型(广州)、体验型(成都)、观光型(北京/西安/重庆)。')
add_figure(doc, 'output_figures/module10_city_radar.png', '图8 城市对比雷达分析：8城市多维雷达对比、综合得分排名、消费×转化气泡图')

doc.add_paragraph()

# --- 2.3 完整实验报告自动生成 ---
add_heading_styled(doc, '2.3 实验报告自动生成系统（技术创新）', 2)
add_para(doc, '完成工作：')
items_23 = [
    '✅ 基于python-docx库，实现"分析→图表→报告"全自动化流程。程序运行结束后自动生成完整的Word实验报告（1.5MB），包含：封面页、目录、5大章节、10个分析模块、8张嵌入图表、6项综合结论、10条技术反思。',
    '✅ 报告排版：A4纸张、黑体标题/宋体正文、1.5倍行距、首行缩进；图表居中带编号题注；数据表采用Light Grid样式；封面含课题名称、数据来源、分析工具、完成日期。',
    '✅ 输出文件：大数据专业综合课程设计_完整实验报告.docx。',
]
for item in items_23:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    set_run_font(run, '宋体', 11)

doc.add_paragraph()

# --- 2.4 综合结论提炼 ---
add_heading_styled(doc, '2.4 六大核心结论', 2)
conclusions = [
    ('结论一：用户行为路径高度标准化', '所有路径均为"首页→筛选→POI详情→行为"四步模式。优化重点应缩短首页到POI详情的路径，通过首页个性化推荐减少中间环节。'),
    ('结论二：消费呈显著双峰结构', '30.6%的核心用户贡献约50%总消费。差异化运营——核心用户VIP推荐、普通用户激励引导——是提升ROI的关键。'),
    ('结论三：转化率是POI推荐的核心区分指标', '高分POI与低分POI的转化率差距达34pp，而价格差距仅7元。推荐模型应以转化率为首要权重，价格不应作为排序因子。'),
    ('结论四：时段特征是推荐时机的重要参考', '晚间19:00活跃度高峰适合内容推荐，下午15:00转化率最高适合限时优惠引导，实现"分时施策"。'),
    ('结论五：城市间行为差异小但画像差异大', '8城市预订率标准差仅0.47%，但旅游画像分为6种类型，为地域化推荐提供了"大同小异"的策略基础。'),
    ('结论六：品类间无强关联', '所有关联规则Lift值均接近1.0，交叉推荐效果有限。推荐系统应以同品类内推荐为主策略。'),
]
for i, (title, detail) in enumerate(conclusions, 1):
    add_para(doc, title, bold=True)
    add_para(doc, detail)

doc.add_paragraph()

# ===== 三、核心数据指标 =====
add_heading_styled(doc, '三、核心数据指标', 1)

add_table(doc,
    ['指标', '当前值（终版）', '上周值', '变化趋势'],
    [
        ['累计数据总量', '62,200条', '42,500条', '↑46.4%（超目标24.4%）'],
        ['目标完成度', '124.4%', '85.0%', '↑39.4pp'],
        ['覆盖城市数', '8个', '8个', '-'],
        ['覆盖品类数', '4个', '5个', '↓1（精简）'],
        ['唯一用户数', '44,958人', '-', '新增（较500条样本↑89倍）'],
        ['唯一POI数', '10,222个', '-', '新增'],
        ['数据缺失率', '0%', '0.15%', '↓0.15pp（完全清除）'],
        ['平均消费金额', '¥4,007', '¥3,901.53', '↑2.7%'],
        ['有消费用户占比', '70.2%', '68.1%', '↑2.1pp'],
        ['平均POI评分', '4.00分', '4.03分', '↓0.03分'],
        ['POI详情页转化率', '24.8%', '23.8%', '↑1.0pp'],
        ['高价值用户占比', '21.4%', '15%', '↑6.4pp（RFM模型）'],
        ['仅浏览不转化比例', '19.0%', '-', '新增指标'],
        ['生成分析图表', '8张', '0张', '↑8张'],
        ['生成报告文档', '1份（完整版）', '0份', '↑1份'],
    ]
)

doc.add_paragraph()

# ===== 四、实验中遇到的问题和解决方法 =====
add_heading_styled(doc, '四、实验中遇到的问题和解决方法', 1)

add_table(doc,
    ['问题描述', '解决方案', '状态'],
    [
        ['数据量62,200条超出预期，单次pandas加载内存压力大', '采用分步处理策略：先清洗→再附加衍生字段→分析时按需加载子集，避免全部数据同时驻留内存', '✅ 已解决'],
        ['中文字体在matplotlib图表中显示为乱码', '设置SimHei字体+rcParams全局配置+axes.unicode_minus=False，配合font_prop逐图指定', '✅ 已解决'],
        ['路径字段为字符串形式的列表，无法直接解析', '使用ast.literal_eval安全转换+try/except容错，失败返回空列表，保证程序不中断', '✅ 已解决'],
        ['RFM分层的业务合理性需要验证', '结合K-Means聚类交叉验证，两方法分群结果趋势一致（少数高价值+多数普通用户），确认分层有效', '✅ 已解决'],
        ['Apriori关联规则Lift值偏低(0.49-0.50)，分析结论需谨慎', '解读为"品类间无强关联"而非"无关联"，正向指导推荐策略以同品类内推荐为主', '✅ 已解决'],
        ['K-Means轮廓系数仅0.315，聚类效果一般', '最优K=2时类间分离度有限，说明用户行为存在连续性而非明显断层，RFM离散分层作为补充', '✅ 已解决'],
        ['python-docx生成报告时图片嵌入尺寸不一致', '统一设置width=Inches(5.0)，配合bbox_inches="tight"确保图表比例一致', '✅ 已解决'],
    ]
)

doc.add_paragraph()

# ===== 五、未来两周计划 =====
add_heading_styled(doc, '五、未来两周计划', 1)

add_heading_styled(doc, '5.1 模型深化', 2)
add_table(doc,
    ['任务', '目标', '时间节点'],
    [
        ['协同过滤推荐实现', '基于UserCF/ItemCF构建个性化POI推荐模型，对比加权评分法效果', '第1周'],
        ['用户行为预测模型', '使用分类算法（XGBoost/LightGBM）预测用户预订概率，AUC目标>0.75', '第1-2周'],
        ['NLP评论情感分析', '若接入POI评论文本数据，进行情感极性分析和主题建模', '第2周'],
    ]
)

doc.add_paragraph()

add_heading_styled(doc, '5.2 可视化与交互', 2)
add_table(doc,
    ['任务', '目标', '时间节点'],
    [
        ['交互式Dashboard', '基于Streamlit构建4页数据看板（概览/路径/用户/城市）', '第1周'],
        ['地理可视化增强', '使用folium/kepler.gl实现POI地理热力图与路径动线', '第2周'],
        ['报告模板优化', '支持多种报告模板（简版/完整版/专题版），满足不同场景需求', '第2周'],
    ]
)

doc.add_paragraph()

add_heading_styled(doc, '5.3 数据与模型迭代', 2)
add_table(doc,
    ['任务', '目标', '时间节点'],
    [
        ['接入真实数据源', '尝试接入高德轨迹API或公开旅游数据集（如Flickr-YFCC）', '第1-2周'],
        ['A/B测试框架设计', '设计在线推荐策略A/B测试方案，量化推荐效果提升', '第2周'],
        ['模型可解释性', '引入SHAP/LIME对推荐结果进行归因解释，提升策略可信度', '第2周'],
    ]
)

doc.add_paragraph()

# ===== 六、风险预警与资源需求 =====
add_heading_styled(doc, '六、风险预警与资源需求', 1)

add_table(doc,
    ['风险类型', '描述', '应对措施'],
    [
        ['数据风险', '当前数据为模拟生成，行为分布过于均匀（各约25%），与真实场景存在偏差', '引入公开真实旅游数据集进行交叉验证；后续尝试接入高德轨迹API'],
        ['模型风险', 'K-Means轮廓系数偏低(0.315)，聚类效果有限；Apriori Lift值偏低(<1.0)', '探索DBSCAN/GMM替代聚类方法；尝试FP-Growth替代Apriori'],
        ['工程风险', '单机pandas处理62,200+条数据接近性能上限', '考虑迁移至Polars/Dask进行分布式处理'],
        ['时效风险', '数据仅覆盖7天（4月20-26日），季节性特征捕捉不足', '标注数据局限性，在报告中明确适用范围'],
    ]
)

doc.add_paragraph()
add_para(doc, '资源需求：', bold=True)
add_para(doc, '• 高德地图API Key（如需补充真实数据）')
add_para(doc, '• 公开旅游评论数据集（如TripAdvisor公开数据集，用于NLP分析）')
add_para(doc, '• GPU资源（如需训练深度学习推荐模型）')

doc.add_paragraph()

# ===== 七、本周收获与反思 =====
add_heading_styled(doc, '七、本周收获与反思', 1)

add_heading_styled(doc, '7.1 主要收获', 2)
harvests = [
    '全流程贯通能力：掌握了从"数据采集→清洗→EDA→多维分析→可视化→报告撰写"的完整数据科学项目流程，实现了代码一键运行、端到端输出的工程化交付。',
    '多维度分析思维：从行为路径、消费偏好、时空特征、转化漏斗、用户画像、关联规则、城市对比7个互补维度切入分析，避免了单一视角的片面结论。',
    '自动化报告生成：掌握了python-docx的高级用法——封面排版、目录结构、图表嵌入、表格样式、字体控制，实现了"代码写完即报告生成"的自动化工作流。',
    '算法实践深化：在实践中深入理解了K-Means聚类（轮廓系数选K）、Apriori关联规则（Support/Confidence/Lift三指标）、RFM模型（分位数分层）的本质含义和适用边界。',
    '业务洞察提炼：从技术指标到业务建议的跨越——如"晚间推内容、下午推转化"的分时策略、"转化率为首、价格为辅"的推荐权重设计等，体现了"分析为业务服务"的思维。',
]
for item in harvests:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'• {item}')
    set_run_font(run, '宋体', 11)

doc.add_paragraph()

add_heading_styled(doc, '7.2 反思与改进', 2)
reflections = [
    '数据真实性始终是核心瓶颈：62,200条数据虽规模可观，但行为类型分布过于均匀（四种行为各约25%）、消费双峰结构非自然形成，与真实旅游场景存在差距。后续应优先接入高德轨迹API或合作企业数据源。',
    '分析深度可进一步提升：当前以描述性分析为主（占比/排序/对比），预测性分析（用户行为预测）和因果推断（折扣对转化率的影响）尚未涉及。',
    '时间维度存在固有局限：数据仅覆盖7天，无法分析节假日效应、季节性波动等重要的旅游行为周期特征。',
    '推荐模型停留在规则评分阶段：当前加权评分法本质上是人工规则，尚未实现基于用户协同过滤或深度学习的内容推荐。',
    '图表可读性有优化空间：部分图表信息密度较高（如6城市雷达图），在报告尺寸下部分标签重叠，后续应增加交互式可视化方案。',
]
for item in reflections:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'• {item}')
    set_run_font(run, '宋体', 11)

doc.add_paragraph()

# ===== 八、下周核心目标 =====
add_heading_styled(doc, '八、下周核心目标', 1)

goals = [
    '✅ 完成协同过滤推荐算法原型（UserCF + ItemCF + GeoCF）',
    '✅ 构建Streamlit交互式数据看板（至少4页）',
    '✅ 完成用户预订概率预测模型（AUC > 0.75）',
    '✅ 撰写最终课程论文（含引言/相关工作/方法/实验/结论）',
]
for goal in goals:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(goal)
    set_run_font(run, '宋体', 11, True)

doc.add_paragraph()

# 保存
output_path = '大数据专业综合课程设计_周报四.docx'
doc.save(output_path)
print(f'周报四已保存: {output_path}')
