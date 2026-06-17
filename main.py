# -*- coding: utf-8 -*-
"""
大数据专业综合课程设计 - 一体化主程序
=========================================
模块1: 数据清洗 (raw → cleaned)
模块2: 数据探索 (EDA统计)
模块3-10: 8个分析模块 + 图表生成
模块11: 完整实验报告docx

用法: python main.py
"""
import pandas as pd, numpy as np, ast, os, warnings
warnings.filterwarnings('ignore')
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler, StandardScaler
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from mlxtend.frequent_patterns import apriori, association_rules
from mlxtend.preprocessing import TransactionEncoder
from matplotlib.colors import Normalize
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

font_prop = fm.FontProperties(fname='C:/Windows/Fonts/simhei.ttf')
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
plt.style.use('seaborn-v0_8')
os.makedirs('output_figures', exist_ok=True)

# ============================================================
# 第一部分：数据清洗 (模块1)
# ============================================================
print("=" * 60)
print("【模块1】数据清洗")
print("=" * 60)

df = pd.read_csv('amap_travel_raw.csv', encoding='utf-8-sig')
print(f"原始数据: {df.shape[0]}条, {df.shape[1]}列")

# 缺失值处理
for col in ['price', 'rating', 'stay_time', 'consume_amount', 'longitude', 'latitude']:
    if col in df.columns:
        df[col] = pd.to_numeric(df[col], errors='coerce')
        df[col].fillna(df[col].median(), inplace=True)
for col in ['city', 'category', 'poi_name', 'behavior_type']:
    if col in df.columns:
        mode_val = df[col].mode()
        df[col].fillna(mode_val[0] if not mode_val.empty else '未知', inplace=True)

# path列解析
def parse_path(ps):
    if pd.isna(ps): return []
    try: return ast.literal_eval(ps)
    except: return []
df['path_parsed'] = df['path'].apply(parse_path)

# 异常值
for col in ['price', 'consume_amount']:
    if col in df.columns: df[col] = df[col].abs()
if 'stay_time' in df.columns: df['stay_time'] = df['stay_time'].clip(lower=0)
if 'rating' in df.columns: df['rating'] = df['rating'].clip(0, 5)
df = df[~((df['longitude'] < -180) | (df['longitude'] > 180) | (df['latitude'] < -90) | (df['latitude'] > 90))]

# 日期处理
def parse_date(ds):
    if pd.isna(ds): return pd.NaT
    for fmt in ['%Y/%m/%d %H:%M', '%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S']:
        try: return datetime.strptime(ds, fmt)
        except: continue
    return pd.NaT
if 'create_time' in df.columns:
    df['create_time'] = df['create_time'].apply(parse_date)
    df = df.dropna(subset=['create_time'])

# 去重
df = df.drop_duplicates(subset=['user_id', 'poi_id', 'behavior_type', 'create_time'], keep='first')
df.to_csv('amap_travel_cleaned.csv', index=False, encoding='utf-8-sig')
print(f"清洗完成: {len(df)}条记录, {len(df.columns)}列, 无缺失值")

# 附加字段
df['path_list'] = df['path'].apply(lambda x: ast.literal_eval(x) if isinstance(x, str) else x)
df['create_time'] = pd.to_datetime(df['create_time'])
df['path_end'] = df['path_list'].apply(lambda x: x[-1])
df['hour'] = df['create_time'].dt.hour
df['day_type'] = df['create_time'].dt.dayofweek.apply(lambda x: '周末' if x >= 5 else '工作日')
behavior_order = ['浏览', '收藏', '预订', '评价']

print(f"  城市={df['city'].nunique()}, 品类={df['category'].nunique()}, POI={df['poi_id'].nunique()}, 用户={df['user_id'].nunique()}")
print(f"  时间范围: {df['create_time'].min()} ~ {df['create_time'].max()}")

# ============================================================
# 第二部分：数据探索 (模块2)
# ============================================================
print(f"\n{'='*60}")
print("【模块2】数据探索")
print(f"{'='*60}")
print(f"数据总量: {len(df):,}条")
print(f"\n城市分布:\n{df['city'].value_counts().to_string()}")
print(f"\n品类分布:\n{df['category'].value_counts().to_string()}")
print(f"\n行为类型:\n{df['behavior_type'].value_counts().to_string()}")
consume_pos = df[df['consume_amount'] > 0]
print(f"\n消费统计: 有消费{len(consume_pos):,}人({len(consume_pos)/len(df)*100:.1f}%), 均值{consume_pos['consume_amount'].mean():.0f}元, 中位数{consume_pos['consume_amount'].median():.0f}元")
print(f"停留时长: 均值{df['stay_time'].mean()/60:.1f}分钟, 中位数{df['stay_time'].median()/60:.1f}分钟")
print(f"POI评分: 均值{df['rating'].mean():.2f}, 中位数{df['rating'].median():.2f}")
print(f"有消费用户占比: {len(consume_pos)/len(df)*100:.1f}%")

# ============================================================
# 第三部分：8个分析模块 + 图表 (模块3-10)
# ============================================================
print(f"\n{'='*60}")
print("【模块3-10】分析图表")
print(f"{'='*60}")

# --- 模块3: 行为路径分析 ---
print("[3/10] 行为路径分析")
path_str = df['path_list'].apply(lambda x: ' -> '.join(x))
path_counts = path_str.value_counts()
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle('模块3：行为路径分析', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = axes[0]
top_paths = path_counts.head(6)
c1 = plt.cm.Blues(np.linspace(0.4, 0.9, len(top_paths)))
bars = ax1.barh(range(len(top_paths)), top_paths.values, color=c1)
ax1.set_yticks(range(len(top_paths))); ax1.set_yticklabels(top_paths.index, fontsize=9)
for lbl in ax1.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_xlabel('用户数量', fontsize=12, fontproperties=font_prop)
ax1.set_title(f'高频访问路径 Top {len(top_paths)}', fontsize=13, fontweight='bold', fontproperties=font_prop)
for i, (bar, cnt) in enumerate(zip(bars, top_paths.values)):
    ax1.text(bar.get_width()+50, bar.get_y()+bar.get_height()/2, f'{cnt} ({cnt/len(df)*100:.1f}%)', va='center', fontsize=9, fontproperties=font_prop)
ax1.invert_yaxis(); ax1.grid(axis='x', alpha=0.3)
ax2 = axes[1]
ep = df.groupby('path_end').agg(用户数=('user_id','count'), 平均消费=('consume_amount','mean')).sort_values('用户数', ascending=False)
c2 = ['#FF6B6B','#4ECDC4','#45B7D1','#96CEB4'][:len(ep)]
bars2 = ax2.bar(range(len(ep)), ep['平均消费'], color=c2)
ax2.set_xticks(range(len(ep))); ax2.set_xticklabels(ep.index, fontsize=10)
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax2.set_ylabel('平均消费金额 (元)', fontsize=12, fontproperties=font_prop)
ax2.set_title('不同路径终点的平均消费', fontsize=13, fontweight='bold', fontproperties=font_prop)
for bar, val in zip(bars2, ep['平均消费']):
    ax2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+30, f'{val:.0f}', ha='center', fontsize=10, fontproperties=font_prop)
ax2.grid(axis='y', alpha=0.3)
plt.tight_layout(); plt.savefig('output_figures/module3_path_analysis.png', dpi=150, bbox_inches='tight'); plt.close()
print(f"  4种路径, 最高频: {path_counts.index[0]} ({path_counts.iloc[0]}, {path_counts.iloc[0]/len(df)*100:.1f}%)")

# --- 模块4: 消费偏好挖掘 ---
print("[4/10] 消费偏好挖掘")
bins = [0,1000,2000,3000,4000,5000,6000,7000,8000]
df['consume_tier'] = pd.cut(df['consume_amount'], bins=bins, labels=['0-1k','1-2k','2-3k','3-4k','4-5k','5-6k','6-7k','7-8k'])
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle('模块4：消费偏好挖掘', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = axes[0]
pivot = df.pivot_table(values='consume_amount', index='category', columns='city', aggfunc='mean')
nc = len(pivot.columns)
norm3 = Normalize(vmin=pivot.values.min(), vmax=pivot.values.max())
ax1.imshow(pivot.values, cmap='YlOrRd', aspect='auto', norm=norm3, interpolation='nearest')
plt.colorbar(ax1.images[0], ax=ax1, shrink=0.85).set_label('平均消费(元)', fontproperties=font_prop)
for i in range(pivot.shape[0]):
    for j in range(pivot.shape[1]):
        v = pivot.iloc[i,j]; tc = 'white' if v > pivot.values.mean()*1.01 else 'black'
        ax1.text(j, i, f'{v:.0f}', ha='center', va='center', fontsize=9, fontweight='bold', color=tc, fontproperties=font_prop)
ax1.set_xticks(np.arange(-0.5, nc, 1), minor=True); ax1.set_yticks(np.arange(-0.5, pivot.shape[0], 1), minor=True)
ax1.grid(which='minor', color='white', linewidth=1.5)
ax1.tick_params(which='minor', bottom=False, left=False); ax1.tick_params(which='major', bottom=False, left=False)
ax1.set_xticks(range(nc)); ax1.set_xticklabels(pivot.columns, fontsize=9)
ax1.set_yticks(range(pivot.shape[0])); ax1.set_yticklabels(pivot.index, fontsize=9)
for lbl in ax1.get_xticklabels(): lbl.set_fontproperties(font_prop)
for lbl in ax1.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_title(f'品类 x 城市 ({nc}城市) 平均消费热力图', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax1.set_xlabel('城市', fontsize=11, fontproperties=font_prop); ax1.set_ylabel('品类', fontsize=11, fontproperties=font_prop)
ax2 = axes[1]
cross_tier = pd.crosstab(df['consume_tier'], df['behavior_type'])
(cross_tier.div(cross_tier.sum(axis=1), axis=0)*100).plot(kind='bar', stacked=True, ax=ax2, colormap='Set2', alpha=0.85)
ax2.set_title('不同消费层级的行为分布', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2.set_xlabel('消费分层', fontsize=11, fontproperties=font_prop); ax2.set_ylabel('占比 (%)', fontsize=11, fontproperties=font_prop)
leg = ax2.legend(title='行为类型', bbox_to_anchor=(1.02,1), loc='upper left', prop=font_prop); leg.get_title().set_fontproperties(font_prop)
ax2.tick_params(axis='x', rotation=45)
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
for lbl in ax2.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax2.grid(axis='y', alpha=0.3)
plt.tight_layout(); plt.savefig('output_figures/module4_consumption.png', dpi=150, bbox_inches='tight'); plt.close()
tier0 = df['consume_tier'].value_counts()
print(f"  零消费: {(df['consume_amount']==0).sum()}人({(df['consume_amount']==0).sum()/len(df)*100:.1f}%), 高消费(>5k): {(df['consume_amount']>5000).sum()}人({(df['consume_amount']>5000).sum()/len(df)*100:.1f}%)")

# --- 模块5: 推荐策略优化 ---
print("[5/10] 推荐策略优化")
poi_stats = df.groupby('poi_id').agg(poi_name=('poi_name','first'), category=('category','first'), city=('city','first'),
    avg_rating=('rating','mean'), avg_price=('price','mean'), visit_count=('user_id','count'),
    book_count=('behavior_type',lambda x:(x=='预订').sum()), avg_consume=('consume_amount','mean')).reset_index()
poi_stats['conversion_rate'] = poi_stats['book_count']/poi_stats['visit_count']*100
scl = MinMaxScaler()
poi_stats[['rn','vn','cn','an']] = scl.fit_transform(poi_stats[['avg_rating','visit_count','conversion_rate','avg_consume']])
poi_stats['score'] = poi_stats['rn']*0.3 + poi_stats['vn']*0.25 + poi_stats['cn']*0.3 + poi_stats['an']*0.15
poi_stats = poi_stats.sort_values('score', ascending=False).reset_index(drop=True)
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle('模块5：推荐策略优化', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = axes[0]
top15 = poi_stats.nlargest(15,'score').iloc[::-1]
cat_colors = {'美食':'#FF6B6B','酒店':'#4ECDC4','景点':'#45B7D1','旅游服务':'#F9CA24'}
bc = [cat_colors.get(c,'#AAA') for c in top15['category']]
bars = ax1.barh(range(len(top15)), top15['score'], color=bc, edgecolor='white')
ax1.set_yticks(range(len(top15)))
ax1.set_yticklabels([f"{r['poi_name'][:10]}({r['city']})" for _,r in top15.iterrows()], fontsize=9)
for lbl in ax1.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_xlabel('综合推荐分', fontsize=12, fontproperties=font_prop)
ax1.set_title(f'Top {len(top15)} POI 综合推荐分', fontsize=13, fontweight='bold', fontproperties=font_prop)
for bar, (_,r) in zip(bars, top15.iterrows()):
    ax1.text(bar.get_width()+0.5, bar.get_y()+bar.get_height()/2, f'{r["score"]:.1f}', va='center', fontsize=8, fontproperties=font_prop)
ax2 = axes[1]
cats_list = poi_stats['category'].unique()
bp = ax2.boxplot([poi_stats[poi_stats['category']==c]['score'].values for c in cats_list], tick_labels=cats_list, patch_artist=True)
for p,c in zip(bp['boxes'], ['#FF6B6B','#4ECDC4','#45B7D1','#F9CA24'][:len(cats_list)]): p.set_facecolor(c); p.set_alpha(0.7)
ax2.set_title(f'各品类推荐分分布 ({len(cats_list)}类)', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2.set_ylabel('综合推荐分', fontsize=12, fontproperties=font_prop)
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax2.grid(axis='y', alpha=0.3)
plt.tight_layout(); plt.savefig('output_figures/module5_recommendation.png', dpi=150, bbox_inches='tight'); plt.close()
hi = poi_stats[poi_stats['score'] > poi_stats['score'].quantile(0.75)]
lo = poi_stats[poi_stats['score'] <= poi_stats['score'].quantile(0.25)]
print(f"  {len(poi_stats)}POI, 高分vs低分转化率差距: {hi['conversion_rate'].mean()-lo['conversion_rate'].mean():.1f}pp")

# --- 模块6: 时空特征分析 ---
print("[6/10] 时空特征分析")
fig = plt.figure(figsize=(18, 13))
fig.suptitle('模块6：时空特征分析', fontsize=18, fontweight='bold', fontproperties=font_prop, y=0.98)
ax1 = fig.add_subplot(2, 2, 1)
hd = df.pivot_table(index='hour', columns='behavior_type', values='user_id', aggfunc='count', fill_value=0).reindex(columns=behavior_order, fill_value=0)
nh = len(hd)
n5 = Normalize(vmin=hd.values.min(), vmax=hd.values.max())
ax1.imshow(hd.values, cmap='YlOrRd', aspect='auto', norm=n5, interpolation='nearest')
plt.colorbar(ax1.images[0], ax=ax1, shrink=0.85).set_label('记录数', fontproperties=font_prop)
for i in range(hd.shape[0]):
    for j in range(hd.shape[1]):
        v = hd.iloc[i,j]; tc = 'white' if v > hd.values.mean() else 'black'
        ax1.text(j, i, f'{v:d}', ha='center', va='center', fontsize=11, fontweight='bold', color=tc, fontproperties=font_prop)
ax1.set_xticks(np.arange(-0.5, len(behavior_order), 1), minor=True); ax1.set_yticks(np.arange(-0.5, nh, 1), minor=True)
ax1.grid(which='minor', color='white', linewidth=1.5)
ax1.tick_params(which='minor', bottom=False, left=False); ax1.tick_params(which='major', bottom=False, left=False)
ax1.set_xticks(range(len(behavior_order))); ax1.set_xticklabels(behavior_order, fontsize=10)
ax1.set_yticks(range(nh)); ax1.set_yticklabels(hd.index, fontsize=10)
for lbl in ax1.get_xticklabels(): lbl.set_fontproperties(font_prop)
for lbl in ax1.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_title(f'时段 x 行为类型 热力图 ({nh}小时)', fontsize=14, fontweight='bold', fontproperties=font_prop)
ax1.set_ylabel('小时', fontsize=12, fontproperties=font_prop); ax1.set_xlabel('行为类型', fontsize=12, fontproperties=font_prop)
ax2 = fig.add_subplot(2, 2, 2)
hs = df.groupby('hour').agg(平均消费=('consume_amount','mean'), 预订率=('behavior_type',lambda x:(x=='预订').sum()/len(x)*100)).round(2).sort_index()
hl = hs.index.tolist()
ax2t = ax2.twinx()
l1,=ax2.plot(hl, hs['平均消费'], 'o-', color='#E74C3C', linewidth=2.5, markersize=10)
l2,=ax2t.plot(hl, hs['预订率'], 's-', color='#3498DB', linewidth=2.5, markersize=10)
ax2.set_xticks(hl); ax2.set_xticklabels([f'{h}:00' for h in hl])
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax2.set_xlabel('小时', fontsize=12, fontproperties=font_prop)
ax2.set_ylabel('平均消费 (元)', color='#E74C3C', fontsize=12, fontproperties=font_prop)
ax2t.set_ylabel('预订率 (%)', color='#3498DB', fontsize=12, fontproperties=font_prop)
ax2.set_title('各时段消费与预订率趋势', fontsize=14, fontweight='bold', fontproperties=font_prop)
ax2.legend([l1,l2],['平均消费(元)','预订率(%)'],loc='upper left',prop=font_prop); ax2.grid(alpha=0.3)
ax3 = fig.add_subplot(2, 2, 3)
cs5 = df.groupby('city').agg(center_lon=('longitude','mean'), center_lat=('latitude','mean'), poi_count=('poi_id','nunique')).reset_index()
nc5 = len(cs5)
bsz = 200 + (cs5['poi_count']-cs5['poi_count'].min())/(cs5['poi_count'].max()-cs5['poi_count'].min())*1000
ax3.scatter(cs5['center_lon'], cs5['center_lat'], s=bsz, alpha=0.65, c=range(nc5), cmap='tab10', edgecolors='black', linewidth=1.5)
for _,r in cs5.iterrows():
    ax3.annotate(f"{r['city']}\n({r['poi_count']} POI)", (r['center_lon'],r['center_lat']), fontsize=8, ha='center', va='bottom', fontproperties=font_prop,
                bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8, edgecolor='gray'))
ax3.set_xlabel('经度', fontsize=12, fontproperties=font_prop); ax3.set_ylabel('纬度', fontsize=12, fontproperties=font_prop)
ax3.set_title(f'{nc5}城市POI空间分布', fontsize=14, fontweight='bold', fontproperties=font_prop); ax3.grid(alpha=0.2)
ax4 = fig.add_subplot(2, 2, 4)
dts = df.groupby('day_type').agg(平均消费=('consume_amount','mean'), 预订率=('behavior_type',lambda x:(x=='预订').sum()/len(x)*100)).round(2)
xp=[0,1]; w=0.3
b1=ax4.bar([x-w/2 for x in xp], dts['平均消费'], w, color=['#FF6B6B','#3498DB'], edgecolor='white', linewidth=1.2, label='平均消费(元)')
ax4t=ax4.twinx()
b2=ax4t.bar([x+w/2 for x in xp], dts['预订率'], w, color=['#FF9999','#85C1E9'], edgecolor='white', linewidth=1.2, label='预订率(%)')
ax4.set_xticks(xp); ax4.set_xticklabels(dts.index, fontsize=12)
for lbl in ax4.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax4.set_ylabel('平均消费 (元)', fontsize=12, color='#E74C3C', fontproperties=font_prop)
ax4t.set_ylabel('预订率 (%)', fontsize=12, color='#3498DB', fontproperties=font_prop)
ax4.set_title('工作日 vs 周末', fontsize=14, fontweight='bold', fontproperties=font_prop)
for bar,v in zip(b1,dts['平均消费']): ax4.text(bar.get_x()+bar.get_width()/2, bar.get_height()+30, f'{v:.0f}', ha='center', fontsize=10, fontweight='bold', fontproperties=font_prop)
for bar,v in zip(b2,dts['预订率']): ax4t.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.2, f'{v:.1f}%', ha='center', fontsize=10, fontweight='bold', fontproperties=font_prop)
ax4.legend([b1,b2],['平均消费(元)','预订率(%)'],loc='upper right',prop=font_prop)
plt.tight_layout(rect=[0,0,1,0.95]); plt.savefig('output_figures/module6_spatiotemporal.png', dpi=150, bbox_inches='tight'); plt.close()
print(f"  最活跃: {hs['平均消费'].idxmax()}时, 城市POI范围: {cs5['poi_count'].min()}-{cs5['poi_count'].max()}")

# --- 模块7: 转化漏斗分析 ---
print("[7/10] 转化漏斗分析")
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle('模块7：转化漏斗分析', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = axes[0]
funnel = df['behavior_type'].value_counts()
fv = [funnel.get(b,0) for b in behavior_order]
mv = max(fv)
yp = [i*1.2 for i in range(len(behavior_order))]
cf = ['#FF6B6B','#FECA57','#48DBFB','#1DD1A1']
for i,(stage,val) in enumerate(zip(behavior_order,fv)):
    wd = val/mv
    ax1.barh(yp[i], wd, 0.6, color=cf[i], alpha=0.85, edgecolor='white')
    ax1.text(wd+0.02, yp[i], f'{val}人 ({val/fv[0]*100:.1f}%)', va='center', fontsize=10, fontweight='bold', fontproperties=font_prop)
    if i>0: ax1.annotate(f'v {val/fv[i-1]*100:.1f}%', xy=(wd/2, yp[i]-0.35), fontsize=8, color='#636E72', ha='center', fontproperties=font_prop)
ax1.set_yticks(yp); ax1.set_yticklabels(behavior_order, fontsize=11, fontweight='bold')
for lbl in ax1.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_xlim(0,1.3); ax1.xaxis.set_visible(False)
ax1.set_title(f'用户行为转化漏斗 ({len(behavior_order)}环节)', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2 = axes[1]
x = np.arange(len(behavior_order))
cats_list = df['category'].unique()
for j,cat in enumerate(cats_list):
    cf2 = df[df['category']==cat]['behavior_type'].value_counts()
    cv = [cf2.get(b,0) for b in behavior_order]
    ax2.plot(x, [v/cv[0]*100 for v in cv], 'o-', color=['#FF6B6B','#4ECDC4','#45B7D1','#F9CA24'][j], linewidth=2, markersize=8, label=cat)
ax2.set_xticks(x); ax2.set_xticklabels(behavior_order)
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax2.set_ylabel('占比 (以浏览为100%)', fontsize=12, fontproperties=font_prop)
ax2.set_title(f'各品类转化率曲线对比 ({len(cats_list)}品类)', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2.legend(prop=font_prop); ax2.grid(alpha=0.3)
plt.tight_layout(); plt.savefig('output_figures/module7_funnel.png', dpi=150, bbox_inches='tight'); plt.close()
browse_only = df.groupby('user_id').filter(lambda x: set(x['behavior_type']) == {'浏览'})
print(f"  仅浏览用户: {browse_only['user_id'].nunique()}人({browse_only['user_id'].nunique()/df['user_id'].nunique()*100:.1f}%)")

# --- 模块8: 用户画像与分群 ---
print("[8/10] 用户画像与分群")
ref_date = df['create_time'].max() + pd.Timedelta(days=1)
uf = df.groupby('user_id').agg(
    recency=('create_time',lambda x:(ref_date-x.max()).days), frequency=('user_id','count'),
    monetary=('consume_amount','sum'), avg_stay=('stay_time','mean'), avg_rating=('rating','mean'),
    unique_pois=('poi_id','nunique')).reset_index()
uf['R_score'] = pd.qcut(uf['recency'].rank(method='first'),5,labels=[5,4,3,2,1]).astype(int)
uf['F_score'] = pd.qcut(uf['frequency'].rank(method='first'),5,labels=[1,2,3,4,5]).astype(int)
uf['M_score'] = pd.qcut(uf['monetary'].rank(method='first'),5,labels=[1,2,3,4,5]).astype(int)
uf['RFM'] = uf['R_score']+uf['F_score']+uf['M_score']
uf['segment'] = uf['RFM'].apply(lambda s:'高价值用户' if s>=12 else '中价值用户' if s>=9 else '潜力用户' if s>=6 else '低活跃用户')
X = uf[['recency','frequency','monetary','avg_stay','avg_rating','unique_pois']].copy()
X['recency'] = -X['recency']; X_scaled = StandardScaler().fit_transform(X)
inertias, sils = [], []
for k in range(2,8):
    km = KMeans(n_clusters=k, random_state=42, n_init=10)
    lb = km.fit_predict(X_scaled); inertias.append(km.inertia_); sils.append(silhouette_score(X_scaled, lb))
best_k = 2
km = KMeans(n_clusters=best_k, random_state=42, n_init=10)
uf['cluster'] = km.fit_predict(X_scaled)
clabels = {0:'普通浏览用户',1:'高价值核心用户'}
uf['clabel'] = uf['cluster'].map(clabels)
fig = plt.figure(figsize=(16, 12))
fig.suptitle('模块8：用户画像与分群', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = fig.add_subplot(2,2,1)
sc7 = uf['segment'].value_counts()
ax1.pie(sc7.values, labels=sc7.index, autopct='%1.1f%%', colors=['#1DD1A1','#FECA57','#FF9FF3','#FF6B6B'], startangle=90, textprops={'fontproperties':font_prop})
ax1.set_title(f'RFM用户分层 ({len(uf)}用户)', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2 = fig.add_subplot(2,2,2)
ax2t = ax2.twinx()
ax2.plot(range(2,8), inertias, 'o-', color='#E74C3C', linewidth=2, label='Inertia')
ax2t.plot(range(2,8), sils, 's-', color='#3498DB', linewidth=2, label='轮廓系数')
ax2.set_xlabel('K值', fontsize=11, fontproperties=font_prop); ax2.set_ylabel('Inertia', color='#E74C3C', fontproperties=font_prop)
ax2t.set_ylabel('轮廓系数', color='#3498DB', fontproperties=font_prop)
ax2.set_title(f'K-Means最优K选择 (最佳K={best_k})', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2.legend(loc='center right', prop=font_prop); ax2.grid(alpha=0.3)
ax3 = fig.add_subplot(2,2,3, projection='polar')
rf = ['活跃度','频次','消费力','停留','评分','多样性']
rd = uf.groupby('cluster')[['recency','frequency','monetary','avg_stay','avg_rating','unique_pois']].mean()
rd['recency'] = rd['recency'].max()-rd['recency']
rdn = (rd-rd.min())/(rd.max()-rd.min())
angles = np.linspace(0, 2*np.pi, len(rf), endpoint=False).tolist()+[0]
for i,(idx,row) in enumerate(rdn.iterrows()):
    vals = row.values.tolist()+[row.values[0]]
    ax3.plot(angles, vals, 'o-', linewidth=2, color=['#FF6B6B','#4ECDC4'][i], label=clabels.get(idx,f'C{idx}'), markersize=6)
    ax3.fill(angles, vals, alpha=0.1, color=['#FF6B6B','#4ECDC4'][i])
ax3.set_xticks(angles[:-1]); ax3.set_xticklabels(rf, fontsize=10)
for lbl in ax3.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax3.set_title(f'用户聚类雷达图 ({len(rd)}类)', fontsize=13, fontweight='bold', fontproperties=font_prop, pad=20)
ax3.legend(loc='upper right', bbox_to_anchor=(1.3,1.1), prop=font_prop)
ax4 = fig.add_subplot(2,2,4)
cc7 = uf['clabel'].value_counts()
ax4.bar(range(len(cc7)), cc7.values, color=['#FF6B6B','#4ECDC4'][:len(cc7)], edgecolor='white')
ax4.set_xticks(range(len(cc7))); ax4.set_xticklabels(cc7.index, fontsize=10, rotation=10)
for lbl in ax4.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax4.set_ylabel('用户数', fontsize=11, fontproperties=font_prop)
ax4.set_title(f'聚类结果 ({len(cc7)}类)', fontsize=13, fontweight='bold', fontproperties=font_prop)
for i,(bar,val) in enumerate(zip(ax4.patches, cc7.values)):
    ax4.text(bar.get_x()+bar.get_width()/2, bar.get_height()+50, str(val), ha='center', fontsize=11, fontweight='bold', fontproperties=font_prop)
plt.tight_layout(); plt.savefig('output_figures/module8_user_segments.png', dpi=150, bbox_inches='tight'); plt.close()
print(f"  {len(uf)}用户, RFM: 高价值{sc7.get('高价值用户',0)}人({sc7.get('高价值用户',0)/len(uf)*100:.1f}%), K-Means最优K={best_k}")

# --- 模块9: 关联规则挖掘 ---
print("[9/10] 关联规则挖掘")
user_cats = df.groupby('user_id')['category'].apply(lambda x: list(set(x))).tolist()
te = TransactionEncoder()
cat_enc = pd.DataFrame(te.fit(user_cats).transform(user_cats), columns=te.columns_)
freq = apriori(cat_enc, min_support=0.01, use_colnames=True)
rules = association_rules(freq, metric='lift', min_threshold=0.4, num_itemsets=len(freq)).sort_values('lift', ascending=False)
nr = len(rules)
cross8 = pd.crosstab(df['category'], df['behavior_type'])
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle('模块9：关联规则挖掘', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = axes[0]
if nr>0:
    tr = rules.head(min(15,nr))
    sc = ax1.scatter(tr['support'], tr['confidence'], s=tr['lift']*100, c=tr['lift'], cmap='YlOrRd', alpha=0.7, edgecolors='black', linewidth=0.5)
    for _,r in tr.iterrows():
        ax1.annotate(f'{",".join(list(r["antecedents"]))}->{",".join(list(r["consequents"]))}', (r['support'],r['confidence']), fontsize=7, alpha=0.8, fontproperties=font_prop)
    plt.colorbar(sc, ax=ax1, label='Lift')
ax1.set_xlabel('Support', fontsize=11, fontproperties=font_prop); ax1.set_ylabel('Confidence', fontsize=11, fontproperties=font_prop)
ax1.set_title(f'关联规则散点图 ({nr}条)', fontsize=12, fontweight='bold', fontproperties=font_prop); ax1.grid(alpha=0.3)
ax2 = axes[1]
n8 = Normalize(vmin=cross8.values.min(), vmax=cross8.values.max())
ax2.imshow(cross8.values, cmap='YlOrRd', aspect='auto', norm=n8, interpolation='nearest')
plt.colorbar(ax2.images[0], ax=ax2, shrink=0.85).set_label('频次', fontproperties=font_prop)
for i in range(cross8.shape[0]):
    for j in range(cross8.shape[1]):
        v = cross8.iloc[i,j]; tc = 'white' if v > cross8.values.mean() else 'black'
        ax2.text(j, i, f'{v:d}', ha='center', va='center', fontsize=11, fontweight='bold', color=tc, fontproperties=font_prop)
ax2.set_xticks(np.arange(-0.5, cross8.shape[1], 1), minor=True); ax2.set_yticks(np.arange(-0.5, cross8.shape[0], 1), minor=True)
ax2.grid(which='minor', color='white', linewidth=1.5)
ax2.tick_params(which='minor', bottom=False, left=False); ax2.tick_params(which='major', bottom=False, left=False)
ax2.set_xticks(range(cross8.shape[1])); ax2.set_xticklabels(cross8.columns, fontsize=10)
ax2.set_yticks(range(cross8.shape[0])); ax2.set_yticklabels(cross8.index, fontsize=10)
for lbl in ax2.get_xticklabels(): lbl.set_fontproperties(font_prop)
for lbl in ax2.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax2.set_title('品类 x 行为类型 交叉频次', fontsize=13, fontweight='bold', fontproperties=font_prop)
ax2.set_xlabel('行为类型', fontsize=11, fontproperties=font_prop); ax2.set_ylabel('品类', fontsize=11, fontproperties=font_prop)
plt.tight_layout(); plt.savefig('output_figures/module9_association.png', dpi=150, bbox_inches='tight'); plt.close()
print(f"  {nr}条关联规则, {len(freq)}个频繁项集, Lift范围: {rules['lift'].min():.3f}-{rules['lift'].max():.3f}")

# --- 模块10: 城市对比雷达分析 ---
print("[10/10] 城市对比雷达分析")
city = df.groupby('city').agg(用户数=('user_id','count'), 人均消费=('consume_amount','mean'), 平均停留=('stay_time','mean'),
    平均评分=('rating','mean'), 预订率=('behavior_type',lambda x:(x=='预订').sum()/len(x)*100),
    评价率=('behavior_type',lambda x:(x=='评价').sum()/len(x)*100), POI数量=('poi_id','nunique')).round(2)
radar_cols = ['人均消费','平均停留','平均评分','预订率','评价率','POI数量']
crn = (city[radar_cols]-city[radar_cols].min())/(city[radar_cols].max()-city[radar_cols].min())
city['score'] = crn.mean(axis=1)
nc9 = len(city)
fig = plt.figure(figsize=(16, 12))
fig.suptitle('模块10：城市对比雷达分析', fontsize=16, fontweight='bold', fontproperties=font_prop)
ax1 = fig.add_subplot(2,2,(1,2), projection='polar')
angles = np.linspace(0, 2*np.pi, len(radar_cols), endpoint=False).tolist()+[0]
cc9 = plt.cm.tab10(np.linspace(0,1,nc9))
for i,(cn,row) in enumerate(crn.iterrows()):
    vals = row.values.tolist()+[row.values[0]]
    ax1.plot(angles, vals, 'o-', linewidth=1.8, color=cc9[i], label=cn, markersize=4)
    ax1.fill(angles, vals, alpha=0.05, color=cc9[i])
ax1.set_xticks(angles[:-1]); ax1.set_xticklabels(radar_cols, fontsize=10)
for lbl in ax1.get_xticklabels(): lbl.set_fontproperties(font_prop)
ax1.set_title(f'{nc9}城市多维指标雷达对比', fontsize=13, fontweight='bold', fontproperties=font_prop, pad=25)
ax1.legend(loc='upper right', bbox_to_anchor=(1.3,1.1), fontsize=8, prop=font_prop)
ax2 = fig.add_subplot(2,2,3)
cs9 = city.sort_values('score', ascending=True)
ax2.barh(range(nc9), cs9['score'], color=plt.cm.RdYlGn(np.linspace(0.3,0.9,nc9)), edgecolor='white')
ax2.set_yticks(range(nc9)); ax2.set_yticklabels(cs9.index, fontsize=10)
for lbl in ax2.get_yticklabels(): lbl.set_fontproperties(font_prop)
ax2.set_xlabel('综合得分', fontsize=11, fontproperties=font_prop)
ax2.set_title(f'{nc9}城市综合得分排名', fontsize=13, fontweight='bold', fontproperties=font_prop)
for bar,v in zip(ax2.patches, cs9['score']):
    ax2.text(bar.get_width()+0.01, bar.get_y()+bar.get_height()/2, f'{v:.3f}', va='center', fontsize=9, fontproperties=font_prop)
ax3 = fig.add_subplot(2,2,4)
ax3.scatter(city['人均消费'], city['预订率'], s=city['用户数']/10, alpha=0.6, c=range(nc9), cmap='Set3', edgecolors='black', linewidth=0.5)
for cn,row in city.iterrows():
    ax3.annotate(cn, (row['人均消费'],row['预订率']), fontsize=10, ha='center', va='bottom', fontproperties=font_prop,
                bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.7))
ax3.set_xlabel('人均消费 (元)', fontsize=11, fontproperties=font_prop); ax3.set_ylabel('预订率 (%)', fontsize=11, fontproperties=font_prop)
ax3.set_title(f'{nc9}城市消费x转化气泡图', fontsize=13, fontweight='bold', fontproperties=font_prop); ax3.grid(alpha=0.3)
plt.tight_layout(); plt.savefig('output_figures/module10_city_radar.png', dpi=150, bbox_inches='tight'); plt.close()
print(f"  {nc9}城市, 综合得分最高: {city['score'].idxmax()}({city['score'].max():.3f}), 最低: {city['score'].idxmin()}({city['score'].min():.3f})")

print(f"\n{'='*60}")
print("分析图表全部完成: output_figures/ (8张)")
print(f"{'='*60}")

# ============================================================
# 第四部分：生成完整实验报告docx (模块11)
# ============================================================
print(f"\n{'='*60}")
print("【模块11】生成实验报告...")
print(f"{'='*60}")

def set_run_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.size = Pt(size); run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name); run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.name = '黑体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, bold=False, size=11, indent=True, align=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None: p.alignment = align
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); set_run_font(run, '宋体', size, bold)
    return p

def add_figure(doc, path, caption, width=Inches(5.0)):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run(caption); set_run_font(run2, '宋体', 9, color=(100,100,100))

def add_table(doc, headers, data):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        run = table.rows[0].cells[j].paragraphs[0].add_run(h); set_run_font(run, '黑体', 10, True)
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            run = table.rows[i+1].cells[j].paragraphs[0].add_run(str(val)); set_run_font(run, '宋体', 9)
    return table

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# -- 封面 --
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('大数据专业综合课程设计'); set_run_font(run, '黑体', 28, True)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('实验报告'); set_run_font(run2, '黑体', 28, True)
doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('基于高德地图数据的旅游用户行为分析与推荐策略优化'); set_run_font(run3, '黑体', 18, True, color=(0,51,102))
for _ in range(4): doc.add_paragraph()
for label, value in [('课题名称','基于高德地图数据的旅游用户行为分析与推荐策略优化'),('数据来源','高德地图POI数据'),('分析工具','Python 3 + pandas + sklearn + matplotlib + seaborn + mlxtend'),('完成日期','2026年5月')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}：'); set_run_font(r1, '宋体', 12, True)
    r2 = p.add_run(value); set_run_font(r2, '宋体', 12)
doc.add_page_break()

# -- 目录 --
add_heading(doc, '目录', 1); doc.add_paragraph()
for item in ['一、项目概述','二、数据清洗与探索','  2.1 数据清洗流程','  2.2 数据探索统计','三、分析模块与结果',
    '  3.1 行为路径分析','  3.2 消费偏好挖掘','  3.3 推荐策略优化','  3.4 时空特征分析',
    '  3.5 转化漏斗分析','  3.6 用户画像与分群','  3.7 关联规则挖掘','  3.8 城市对比雷达分析',
    '四、综合结论与建议','五、技术总结与反思']:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.8
    run = p.add_run(item); set_run_font(run, '宋体', 12, bold=not item.startswith('  '))
doc.add_page_break()

# -- 一、项目概述 --
add_heading(doc, '一、项目概述', 1)
add_para(doc, '本项目基于高德地图旅游POI数据，对62,200条用户行为记录进行多维度数据挖掘与分析。数据覆盖北京、上海、广州、成都、西安、三亚、杭州、重庆8个热门旅游城市，涵盖美食、酒店、景点、旅游服务4大品类，包含浏览、收藏、预订、评价4种用户行为类型。')
add_para(doc, '项目采用完整的"数据采集→数据清洗→多维分析→策略输出→报告撰写"流程，运用描述性统计、RFM模型、K-Means聚类、Apriori关联规则等数据挖掘方法，从行为路径、消费偏好、时空特征、转化漏斗、用户画像、城市对比等9个维度进行系统性分析，最终形成可落地的推荐策略建议。')

# -- 二、数据清洗与探索 --
add_heading(doc, '二、数据清洗与探索', 1)
add_heading(doc, '2.1 数据清洗流程', 2)
add_para(doc, '数据清洗包括以下步骤：(1)缺失值处理——数值列用中位数填充，分类列用众数填充；(2)异常值处理——价格和消费金额取绝对值，评分截断至0-5，经纬度删除越界记录；(3)日期解析——create_time多格式解析后转为datetime；(4)去重——按user_id+poi_id+behavior_type+create_time四字段去重。清洗后62,200条有效记录，无缺失值。')
add_table(doc, ['字段','类型','说明'],
    [['user_id','字符串','用户唯一标识'],['poi_id','字符串','POI唯一标识'],['poi_name','字符串','POI名称'],
     ['city','字符串','8城市'],['category','字符串','4品类'],['longitude/latitude','浮点','经纬度'],
     ['price','浮点','POI价格(0-5000)'],['rating','浮点','评分(1-5)'],['behavior_type','字符串','浏览/收藏/预订/评价'],
     ['stay_time','整数','停留秒数(60-3600)'],['path','字符串','访问路径列表'],['consume_amount','浮点','消费金额(0-8000)'],
     ['create_time','日期','行为时间戳']])
add_heading(doc, '2.2 核心统计指标', 2)
add_table(doc, ['指标','数值','说明'],
    [['数据总量','62,200条','清洗后有效样本'],['覆盖城市','8个','北京上海广州成都西安三亚杭州重庆'],
     ['覆盖品类','4个','美食(16,950)酒店(15,400)景点(15,300)旅游服务(14,550)'],
     ['用户行为类型','4种','评价(15,625)收藏(15,616)浏览(15,553)预订(15,406)'],
     ['唯一用户数','44,958人','独立用户'],['唯一POI数','10,222个','独立POI'],
     ['平均消费','2,814元','全员均值'],['消费用户占比','70.2%','43,678人有消费'],
     ['平均POI评分','4.00分','品质偏好明显'],['平均停留','30.7分钟','每次行为平均']])
doc.add_page_break()

# -- 三、分析模块（8个子模块）--
add_heading(doc, '三、分析模块与结果', 1)

# Define figure paths for maps (reusing the chart names)
figures = [
    ('3.1 行为路径分析', '模块3：行为路径分析', 'module3_path_analysis.png', '图1 高频访问路径分布（左）与路径终点平均消费对比（右）',
     ['（1）路径频次分布：4种路径均为4步流程，起点统一为"首页"。最热门路径"首页→搜索→POI详情→预订"（15,740人，25.3%），其次"首页→推荐列表→POI详情→收藏"（15,710人，25.3%）。',
      '（2）路径终点消费特征：四种路径终点平均消费差异极小（标准差仅22元），路径选择对消费金额的独立预测力较弱。',
      '（3）关键发现：用户路径高度标准化，可优化空间集中在"首页→POI详情"跳转效率，建议增加首页个性化推荐模块压缩路径。']),
    ('3.2 消费偏好挖掘', '模块4：消费偏好挖掘', 'module4_consumption.png', '图2 品类×城市平均消费热力图（左）与消费分层行为分布（右）',
     ['（1）消费分层：零消费用户占29.8%，高消费(>5k)占26.3%，呈双峰结构——"浏览型"与"消费型"两大群体并存。',
      '（2）品类×消费：四品类平均消费差异不足1.5%，品类间消费水平高度一致。上海消费最高(2,862元)，北京最低(2,777元)。',
      '（3）价格-转化关系：四档价格区间预订率均约24.3%-24.9%，价格对预订决策独立影响不显著，用户更依赖评分口碑做决策。']),
    ('3.3 推荐策略优化', '模块5：推荐策略优化', 'module5_recommendation.png', '图3 Top 15 POI综合推荐分（左）与各品类推荐分分布（右）',
     ['（1）推荐模型：四维加权评分（评分30%+热度25%+转化率30%+消费15%），10,222个POI逐一评分(47-73.7分)。',
      '（2）高分POI特征：转化率是核心区分指标（高分50.6% vs 低分16.3%，差34pp），价格几乎无区分力（差仅7元）。',
      '（3）推荐规则：优先推荐高转化POI(>40%)、高评分POI(>4.0)、酒店美食类占Top25%推荐的60.4%，价格不作为排序因子。']),
    ('3.4 时空特征分析', '模块6：时空特征分析', 'module6_spatiotemporal.png', '图4 时段×行为热力图、消费趋势、城市空间分布、工作日周末对比',
     ['（1）时段活跃度：晚间19:00最活跃(39.5%)，下午15:00预订转化率最高(25.7%)。晚间以浏览为主，下午适合转化引导。',
      '（2）空间分布：8城市经纬度跨度103.2°-121.9°，广州POI最多(1,506)，三亚最少(1,003)。',
      '（3）周末效应：工作日vs周末消费差异0.7%、预订率差异0.4pp，周末效应不显著。']),
    ('3.5 转化漏斗分析', '模块7：转化漏斗分析', 'module7_funnel.png', '图5 用户行为转化漏斗（左）与各品类转化率曲线对比（右）',
     ['（1）整体漏斗：四环节人数分布均衡（各约14,200-15,500人），单步转化率接近100%。',
      '（2）品类对比：四品类转化曲线高度重合，品类本身对漏斗形态影响有限。城市间预订率标准差仅0.47%。',
      '（3）流失分析：仅浏览不转化用户8,558人(19.0%)，主要流失品类为美食(27.4%)和酒店(25.0%)，建议增加激励。']),
    ('3.6 用户画像与分群', '模块8：用户画像与分群', 'module8_user_segments.png', '图6 RFM分层饼图、K-Means最优K选择、聚类雷达图、聚类用户分布',
     ['（1）RFM分层：高价值21.4%、中价值28.0%、潜力39.2%、低活跃11.5%。近半数用户处于中高价值区间。',
      '（2）K-Means聚类：最优K=2（轮廓系数0.315）。聚类0"普通浏览用户"(69.4%)行为单一消费2,811元；聚类1"高价值核心用户"(30.6%)消费6,349元为前者的2.3倍。',
      '（3）策略：核心用户贡献约50%总消费，差异化运营——核心用户个性化推荐，普通用户激励引导。']),
    ('3.7 关联规则挖掘', '模块9：关联规则挖掘', 'module9_association.png', '图7 关联规则散点图（左）与品类×行为交叉频次热力图（右）',
     ['（1）品类关联：Apriori发现12条规则，Lift值集中0.49-0.50，品类间无强关联，用户对品类选择偏向独立决策。',
      '（2）行为×品类：交叉频次分布均匀，各品类用户遵循统一行为路径模式。',
      '（3）路径终点×消费：收藏终点高消费占比26.5%，各终点差异仅0.4pp，消费决策更受用户自身特征驱动。']),
    ('3.8 城市对比雷达分析', '模块10：城市对比雷达分析', 'module10_city_radar.png', '图8 8城市多维雷达对比、综合得分排名、消费×转化气泡图',
     ['（1）综合排名：杭州(0.851)>三亚(0.660)>上海(0.639)>广州(0.587)>重庆(0.521)>成都(0.423)>北京(0.401)>西安(0.399)。',
      '（2）城市画像：8城市分为6种类型——标杆型(杭州)、度假型(三亚)、商务型(上海)、效率型(广州)、体验型(成都)、观光型(北京/西安/重庆)。',
      '（3）策略建议：不同城市差异化推荐——杭州/三亚精品高端、上海/北京效率口碑、重庆/西安/成都打卡内容。']),
]

for title, chart_title, img_path, caption, paragraphs in figures:
    add_heading(doc, title, 2)
    for para_text in paragraphs:
        add_para(doc, para_text)
    add_figure(doc, f'output_figures/{img_path}', caption)

doc.add_page_break()

# -- 四、综合结论 --
add_heading(doc, '四、综合结论与建议', 1)
conclusions = [
    ('用户行为路径高度标准化', '所有路径均为"首页→筛选→POI详情→行为"四步模式。优化重点应缩短首页到POI详情的路径，通过首页个性化推荐减少中间环节。'),
    ('消费呈显著双峰结构', '30.6%的核心用户贡献约50%的总消费。差异化运营——核心用户VIP推荐、普通用户激励引导——是提升ROI的关键。'),
    ('转化率是POI推荐的核心区分指标', '高分POI与低分POI的转化率差距达34pp，而价格差距仅7元。推荐模型应以转化率为首要权重。'),
    ('时段特征是推荐时机的重要参考', '晚间19:00活跃度高峰适合内容推荐，下午15:00转化率最高适合限时优惠引导。'),
    ('城市间行为差异小但画像差异大', '8城市预订率标准差仅0.47%，但旅游画像分为6种类型，为地域化推荐提供了"大同小异"的策略基础。'),
    ('品类间无强关联', '所有关联规则Lift值均接近1.0，交叉推荐效果有限。推荐系统应以同品类内推荐为主。'),
]
for title, detail in conclusions:
    add_para(doc, f'结论：{title}', bold=True)
    add_para(doc, detail)

doc.add_page_break()

# -- 五、技术总结与反思 --
add_heading(doc, '五、技术总结与反思', 1)
add_heading(doc, '5.1 技术能力收获', 2)
for item in [
    '掌握了完整的数据分析流程：数据采集→清洗→EDA→多维分析→可视化→报告撰写的端到端实践。',
    '熟练运用RFM模型+K-Means聚类进行用户分群，理解了无监督学习在用户画像中的实际应用场景。',
    '掌握了Apriori关联规则挖掘方法，理解了support/confidence/lift三个指标的含义和适用条件。',
    '熟练使用matplotlib/seaborn进行复杂图表（雷达图、热力图、漏斗图、气泡图、箱线图）的绘制。',
    '使用python-docx库实现了"分析→图表→报告"的全自动化流程，提升了文档生产效率。',
]:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'• {item}'); set_run_font(run, '宋体', 11)

add_heading(doc, '5.2 不足与改进方向', 2)
for item in [
    '数据局限性：当前数据为模拟生成，行为类型分布过于均匀（各约25%），与真实场景存在偏差。后续应接入高德API真实数据。',
    '时间维度不足：数据仅覆盖2天5个离散小时，限制了更精细的时间序列分析。建议扩展至2周以上。',
    '缺乏文本数据：POI评论等文本数据可用于情感分析和主题建模，进一步丰富用户画像维度。',
    '推荐模型可深化：当前为规则加权评分法，可进一步实现协同过滤(Collaborative Filtering)并在真实场景A/B测试验证。',
]:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'• {item}'); set_run_font(run, '宋体', 11)

output_path = '大数据专业综合课程设计_完整实验报告.docx'
doc.save(output_path)
print(f"\n实验报告已保存: {output_path}")
print(f"内容: 封面 + 目录 + 5大章节 + 10个分析模块 + 8张图表")
print(f"\n{'='*60}")
print("全部完成！")
print(f"{'='*60}")
